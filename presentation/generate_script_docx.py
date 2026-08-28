"""
SparkLabs 서류 추적기 — Speaker Script Generator (v2)
All 16 slides covered.
Output: script-v2.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x2B, 0x55)
BLUE   = RGBColor(0x1A, 0x56, 0xAA)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
RED    = RGBColor(0xB9, 0x1C, 0x1C)
GREEN  = RGBColor(0x15, 0x80, 0x3D)
TEAL   = RGBColor(0x0D, 0x6E, 0x6E)
BLACK  = RGBColor(0x1E, 0x29, 0x3B)
ORANGE = RGBColor(0xC2, 0x41, 0x0C)

# ── Style helpers ─────────────────────────────────────────────────────────────
def style_run(run, size, bold=False, italic=False, color=BLACK):
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color
    run.font.name      = "Malgun Gothic"

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    style_run(run, 15 if level == 1 else 13, bold=True,
              color=NAVY if level == 1 else BLUE)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1A56AA")
    return p

def korean_with_translation(korean, translation):
    """Red shaded box with 【읽기】 label + teal EN translation line."""
    # Korean sentence
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.3)
    label = p.add_run("【읽기】 ")
    style_run(label, 10, bold=True, color=RED)
    run = p.add_run(korean)
    style_run(run, 12, color=BLACK)
    pPr = p._p.get_or_add_pPr()
    shd = etree.SubElement(pPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF3F3")

    # English translation
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(7)
    p2.paragraph_format.left_indent  = Inches(0.3)
    label2 = p2.add_run("  ↳ EN:  ")
    style_run(label2, 10, bold=True, color=TEAL)
    run2 = p2.add_run(translation)
    style_run(run2, 11, italic=True, color=TEAL)
    pPr2 = p2._p.get_or_add_pPr()
    shd2 = etree.SubElement(pPr2, qn("w:shd"))
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "F0FDFA")

def bullet(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.25)
    run = p.add_run("•  " + text)
    style_run(run, 11, color=BLACK)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("→ " + text)
    style_run(run, 10, italic=True, color=GRAY)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 60)
    style_run(run, 9, color=RGBColor(0xCC, 0xD5, 0xE0))

def slide_label(num, title, time_est):
    divider()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f"SLIDE {num}  ")
    style_run(r1, 11, bold=True, color=BLUE)
    r2 = p.add_run(f"— {title}")
    style_run(r2, 11, bold=True, color=NAVY)
    r3 = p.add_run(f"   ({time_est})")
    style_run(r3, 10, italic=True, color=GRAY)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SparkLabs 서류 추적기")
style_run(r, 22, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("발표 스크립트  |  Presentation Script")
style_run(r, 13, italic=True, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026년 8월  |  AI Agent 인턴십 최종 발표  |  16슬라이드, 약 15~20분")
style_run(r, 11, color=GRAY)

doc.add_paragraph()

# Legend
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.3)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
pPr = p._p.get_or_add_pPr()
shd = etree.SubElement(pPr, qn("w:shd"))
shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "EBF5FF")
r = p.add_run("읽는 법:  ")
style_run(r, 10, bold=True, color=NAVY)
r2 = p.add_run("일반 bullet = 본인 말로 편하게     ")
style_run(r2, 10, color=BLACK)
r3 = p.add_run("【읽기】 빨간 박스 = 연습해서 읽을 한국어     ")
style_run(r3, 10, color=RED)
r4 = p.add_run("↳ EN: 파란 줄 = 그 문장의 영어 번역")
style_run(r4, 10, color=TEAL)

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 1 — TITLE
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(1, "TITLE", "약 30초")
korean_with_translation(
    "안녕하세요. 저는 이번 여름 스파크랙 코리아에서 AI 에이전트 인턴십을 진행한 [이름]입니다. "
    "오늘은 제가 개발한 '서류 추적기'를 소개하겠습니다.",
    "Hello. My name is [Name], and this summer I completed an AI agent internship at SparkLabs Korea. "
    "Today I’d like to introduce the Document Tracker I built."
)
bullet("This is an internal web platform I built for the SparkLabs investment team.")
bullet("It automates the full investment pipeline — from document collection all the way through to SAFE conversion.")
note("자신 있게 시작 — smile, eye contact. 첫 인상이 중요합니다.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 2 — 목차
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(2, "목차 (Table of Contents)", "약 20초")
bullet("Briefly walk through the 8 agenda cards — just read the titles on screen.")
korean_with_translation(
    "오늘 발표는 크게 8개 파트로 나뉘어 있고, 마지막에 실제 사이트를 직접 시연합니다.",
    "Today’s presentation is divided into 8 main parts, and at the end I’ll do a live demo of the actual website."
)
note("No need to explain each item — the audience will see them soon. Keep this under 20 seconds.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 3 — 웹사이트 소개
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(3, "웹사이트 소개", "약 1분")
korean_with_translation(
    "이 플랫폼은 스파크랙 투자팀을 위한 내부 웹 서비스로, "
    "투자 파이프라인 전 과정을 자동화하는 것을 목표로 개발했습니다.",
    "This platform is an internal web service for the SparkLabs investment team, "
    "built to automate the full investment pipeline."
)
bullet("Covers everything: document collection, due diligence, agreements, execution, and SAFE conversion.")
bullet("RCPS = 상환전환우선주  (Redeemable Convertible Preferred Stock)", 1)
bullet("SAFE = 조건부지분인수계약서  (Simple Agreement for Future Equity)", 1)
bullet("Google Gemini API powers 6 different AI features inside the app.")
bullet("Built with Next.js 16, deployed on Vercel — accessible from anywhere with a SparkLabs account.")
note("Point to the tech stack card on the right and the URL.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 4 — 기존 업무의 문제점
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(4, "기존 업무의 문제점", "약 1분 30초")
korean_with_translation(
    "이 플랫폼을 개발하기 전, 투자 계약서 한 건을 작성하는 데 평균 2~3시간이 소요됐습니다.",
    "Before this platform was built, drafting a single investment agreement took an average of 2 to 3 hours."
)
bullet("Everything was manual — open a Word template, fill every blank by hand.")
bullet("Same info (company name, date, amount) typed repeatedly in many places.")
bullet("Documents scattered across emails and folders — no central tracking.")
bullet("No structured due diligence checklist — no way to track item completion.")
bullet("Bank 30-day payment deadline tracked in spreadsheets — risk of missing it.")
bullet("Process questions during work → had to interrupt the mentor for every answer.")
korean_with_translation(
    "특히 같은 정보를 계약서의 여러 곳에 반복 입력해야 한다는 점이 시간 낙비의 가장 큰 원인이었습니다.",
    "In particular, having to type the same information repeatedly in many places was the biggest cause of wasted time."
)
note("Point at the big '2~3시간' number — pause 2–3 seconds and let it land.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 5 — 전체 기능 개요 (Pipeline Map)
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(5, "전체 기능 개요 (Pipeline Map)", "약 1분 30초")
korean_with_translation(
    "이 플랫폼은 투자 프로세스의 전 과정을 6단계 파이프라인으로 구조화했습니다.",
    "This platform structures the full investment process into a 6-stage pipeline."
)
bullet("Stage 1 → Doc Collection (Feature 2):  drag-and-drop upload, AI classification, Drive sync")
bullet("Stage 2 → Due Diligence (Feature 3):  13-point checklist, AI analysis, per-item notes")
bullet("Stage 3 → Agreement Drafting (Feature 4):  RCPS 69 tokens, SAFE 40 tokens, .docx download")
bullet("Stage 4 → Execution Tracking (Feature 5):  bank instruction, 30-day deadline, email draft")
bullet("Stage 5 → SAFE Conversion (Feature 6):  pre/during/post stages, AI-assisted calculation")
bullet("Stage 6 → Complete")
note("Explain the top bar: Dashboard and Action Center sit above all stages.")
note("Explain the bottom bar: AI chatbot is available on every page; full bilingual toggle throughout.")
korean_with_translation(
    "총 11가지 기능이 있고, 그 중 6가지는 Google Gemini API를 활용합니다.",
    "There are 11 features in total, and 6 of them use the Google Gemini API."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 6 — 기능 ① 메인 대시보드 & 액션 센터
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(6, "기능 ①  메인 대시보드 & 액션 센터", "약 1분 30초")
korean_with_translation(
    "메인 대시보드에서는 현재 추적 중인 모든 딕의 현황을 한눈에 확인할 수 있습니다.",
    "The main dashboard lets you see the status of all currently tracked deals at a glance."
)
bullet("4 KPI tiles across the top: companies tracked / docs ready / DD complete / needs attention")
bullet("Per-company progress bars for both document collection and diligence completion")
bullet("Click any deal → goes to the Deal Overview page — a one-screen summary of all 6 stages")

heading("액션 센터 (Action Center)", 2)
korean_with_translation(
    "액션 센터는 '지금 당장 나 뒤 해야 하는 일이 뫐지?'를 알려주는 패널입니다.",
    "The Action Center is the panel that tells you: what do I need to do right now?"
)
bullet("Priority 1: bank 30-day payment deadline — hard cutoff, shown prominently with date")
bullet("Priority 2: missing required documents — lists the deal and which docs are missing")
bullet("Priority 3: outstanding diligence items — lists the deal and how many items remain")
note("This replaces the manual spreadsheet tracking for deadlines. Nothing falls through the cracks.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 7 — 기능 ② 서류 수집 & AI 분류 & Drive 연동
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(7, "기능 ②  서류 수집 & AI 분류 & Drive 연동", "약 1분 30초")
korean_with_translation(
    "서류 수집 탭에서는 계약서나 관련 서류를 웹사이트에 직접 업로드할 수 있습니다.",
    "In the Document Collection tab, you can upload agreements and related documents directly to the website."
)
bullet("Drag-and-drop upload — no folder navigation needed")
bullet("Required document checklist per deal — automatically checks off when the matching file is uploaded")
bullet("Auto-refreshes every 5 seconds so the status stays current without manual reload")
bullet("Files stored privately on Vercel Blob — not publicly accessible")

heading("AI 파일명 자동 분류  (Google Gemini)", 2)
korean_with_translation(
    "파일명이 '최종본_v3.pdf'처럼 비표준적일 때, Gemini가 문서 종류를 자동으로 판별합니다.",
    "When a filename is non-standard — like 'final_v3.pdf' — Gemini automatically identifies what type of document it is."
)
bullet("Example: 'final_v3.pdf' → auto-classified as certificate of registration (등기부등본)")
bullet("Cost: ~$0.0002 per call, ~150 calls/year → ~$0.03/year")

heading("Google Drive 연동", 2)
bullet("One click → deal folder auto-created in Drive")
bullet("Automatically shared with the entire @sparklabs.co.kr team")
bullet("Uploaded files mirrored to the Drive folder automatically")
note("Screenshot placeholder is on the right — you’ll show the real thing in the demo.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 8 — 기능 ③ 서류 실사 체크리스트
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(8, "기능 ③  서류 실사 체크리스트", "약 1분")
korean_with_translation(
    "서류 실사 탭에서는 멘토님의 내부 문서를 기반으로 한 13항목 체크리스트를 활용합니다.",
    "The Due Diligence tab uses a 13-point checklist based on the mentor’s internal process document."
)
bullet("Points 1–7: document verification — certificate of registration, shareholder register, financials, etc.")
bullet("Points 8–13: process steps — confirming payment deposit, registry changes, follow-up actions")
bullet("Each item has a notes/comments field — can record what was checked and when")
bullet("Per-item AI analysis (Gemini) — suggests review points for that specific check")
bullet("Completed items auto-check; overall progress feeds back into the dashboard KPI tiles")
note("This replaces unstructured mental tracking. Every diligence step is now visible and accountable.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 9 — 기능 ④ 투자 계약서 자동 생성
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(9, "기능 ④  투자 계약서 자동 생성", "약 1분 30초")
korean_with_translation(
    "RCPS 계약서 탭에서는 상환전환우선주 투자계약서를 자동으로 생성할 수 있습니다.",
    "In the RCPS tab, you can automatically generate a Redeemable Convertible Preferred Stock investment agreement."
)
bullet("Template has 69 blank fields — all filled automatically from the form")
bullet("Repeated values (date, party names, amounts) entered once → appear everywhere they're needed")
bullet("Covers signature block, Appendix 1 & 2 automatically")
bullet("5 sections: dates / parties / investment terms / notices / standard clauses")
bullet("Standard clauses have SparkLabs defaults already pre-filled — only change what’s different per deal")
bullet("Download as .docx instantly")

korean_with_translation(
    "SAFE 계약서는 조건부지분인수계약서로, 40개의 빈칸을 자동으로 채움니다.",
    "The SAFE agreement — Simple Agreement for Future Equity — automatically fills 40 blank fields."
)
bullet("Includes valuation cap, discount rate, payment deadline")
bullet("Interested party info fills both Appendix 1 and Appendix 3 simultaneously")
bullet("Standard clauses pre-filled (12% penalty, 120% damages…)")
note("This is the slide where you show the before/after time impact — 2–3 hours → ~3 minutes.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 10 — 기능 ⑤ 투자 집행 추적
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(10, "기능 ⑤  투자 집행 추적", "약 1분")
korean_with_translation(
    "투자 집행 탭에서는 투자 집행 후 필수 절차를 단계별로 추적합니다.",
    "The Investment Execution tab tracks required steps after investment execution, step by step."
)
bullet("운용지시 (bank payment instruction) tracking — confirms the instruction was sent")
bullet("Post-payment document collection management — tracks what’s still needed")
bullet("Custodian bank 30-day deadline — hard cutoff, displayed prominently and surfaced in Action Center")
bullet("Fund type selection: Government-backed (모태평드) vs Private fund")
bullet("Investment structure recorded: method, amount, equity percentage")
bullet("Email draft auto-generation (AI-assisted) — creates a draft for the team member to send")
bullet("Auto-save — data is preserved even if the page is closed mid-entry")
note("The 30-day bank deadline is the hardest cutoff in the whole process — this makes it impossible to miss.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 11 — 기능 ⑥ SAFE 전환 추적 + AI 챗봇 어시스턴트
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(11, "기능 ⑥  SAFE 전환 추적 + AI 챗봇 어시스턴트", "약 1분 30초")
korean_with_translation(
    "전환 추적 탭에서는 SAFE 주식 전환 프로세스를 단계별로 추적합니다.",
    "The Conversion Tracker tab tracks the SAFE-to-equity conversion process step by step."
)
bullet("3 stages: Pre-conversion · During conversion · Post-conversion")
bullet("Deadline auto-calculation and D-day countdown")
bullet("AI-assisted share count calculation (Gemini) — based on principal · valuation cap · discount rate")
bullet("Post-conversion registry change tracking")

heading("AI 프로세스 어시스턴트  (우하단 플로팅 챗봇)", 2)
korean_with_translation(
    "그리고 모든 페이지 우하단에는 AI 프로세스 어시스턴트가 떠 있습니다.",
    "And at the bottom-right of every single page, there’s the AI Process Assistant chatbot."
)
bullet("Grounded in SparkLabs internal process documents — answers are accurate to how SparkLabs actually works")
bullet("Context-aware: knows which deal you’re currently viewing")
bullet("Answers in both Korean and English — whichever language you ask in")
bullet("AI field suggestions for agreements, number extraction for execution, conversion calculation")
bullet("Replaces having to interrupt the mentor for every process question")
note("Mention the bilingual footer: the entire UI can toggle between Korean and English instantly.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 12 — 활용한 AI 도구
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(12, "활용한 AI 도구", "약 1분 30초")

heading("AI 활용 ①  앱 내부에서 사용한 AI", 2)
korean_with_translation(
    "앱 내부에서는 Google의 Gemini API를 활용했습니다.",
    "Inside the app itself, I used Google’s Gemini API."
)
bullet("Model: gemini-flash-lite-latest — Google’s cheapest current model")
bullet("6 use cases in the app:")
bullet("Filename auto-classification (Doc Collection tab)", 1)
bullet("Diligence AI analysis per item (Diligence tab)", 1)
bullet("Agreement field AI suggestions (Agreement tab)", 1)
bullet("Execution number auto-extraction (Execution tab)", 1)
bullet("SAFE conversion share count calculation (Conversion tab)", 1)
bullet("Chatbot process assistant (every page)", 1)
bullet("~$0.0002 per call — see next slide for full cost breakdown")

heading("AI 활용 ②  개발에 활용한 AI 도구", 2)
korean_with_translation(
    "그리고 이 플랫폼을 개발하는 데에는 Anthropic의 Claude와 Claude Code를 사용했습니다.",
    "And to develop this platform, I used Anthropic’s Claude and Claude Code."
)
bullet("Claude — code writing, debugging, feature design, DOCX XML manipulation")
bullet("Claude Code — AI coding assistant as a CLI tool in the terminal")
bullet("I could describe what I wanted in plain language; it would read and edit files directly", 1)
bullet("Made it possible to build all 11 features with no prior programming experience", 1)
note("Left panel = Gemini (inside the app). Right panel = Claude/Claude Code (used to build it).")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 13 — Gemini API 비용 분석
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(13, "Google Gemini API — 비용 분석", "약 1분")
bullet("This slide shows the actual real-world cost of the AI features in this app.")
bullet("Model: gemini-flash-lite-latest — Google’s cheapest current model")
bullet("Cost per API call: ~$0.0002  (that’s 2 hundredths of a cent)")
bullet("Three use cases in the app:")
bullet("Filename classification — ~150 calls/year → ~$0.03", 1)
bullet("Diligence AI analysis — ~50 calls/year → ~$0.01", 1)
bullet("Chatbot assistant — ~250 calls/year → ~$0.05", 1)
bullet("Annual total: ~450 calls → ~$0.09/year")
bullet("Even at 10× usage: ~$0.90/year.  Even at 100×: ~$9.00/year.")
bullet("Google’s free tier covers 60 requests/minute — current usage is well within that")
korean_with_translation(
    "AI 기능을 1년 동안 운영하는 데 드는 비용이 스타빅스 커피 한 잔보다 저렴합니다.",
    "The cost of running the AI features for an entire year is less than a single Starbucks coffee."
)
note("Point at the green callout on the right: $0.09 vs ~$5 coffee. Pause and let that land.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 14 — 성과 및 개선 사항
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(14, "성과 및 개선 사항", "약 1분")
note("Pause 3–4 seconds after this slide appears — let the numbers hit first.")
korean_with_translation(
    "기존에 2~3시간이 걸리던 계약서 작성이, 이제 약 3분으로 줄었습니다. 약 97% 이상의 시간이 단축된 것입니다.",
    "Agreement drafting that used to take 2 to 3 hours now takes about 3 minutes — a reduction of over 97%."
)
bullet("Multiple deals per month — this adds up very fast across the team.")
bullet("Beyond time savings:")
bullet("Bank deadlines surfaced automatically — never missed by accident", 1)
bullet("All documents tracked in one place — no more searching through emails", 1)
bullet("Diligence structured and accountable — every item tracked per deal", 1)
bullet("AI answers process questions instantly — mentor no longer needed for every question", 1)
bullet("Full bilingual support — usable by Korean and English speakers alike", 1)
note("Point at the green '~3분' box — let it sink in.")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 15 — 한계점 및 향후 과제
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(15, "한계점 및 향후 과제", "약 1분 30초")
korean_with_translation(
    "아직 개선해야 할 점도 많습니다.",
    "There are still many areas that need improvement."
)

heading("현재 한계점", 2)
bullet("Template updates are manual — when legal language changes, XML must be edited directly")
bullet("Multiple interested parties not yet supported (currently only 1 per deal)")
bullet("Mobile not fully optimized — desktop-first for now")
bullet("No e-signature — wet signatures still required after download")
bullet("No user permission management — all team members see all deals")
bullet("Diligence checklist item count is hard to adjust without a code change")

heading("앞으로 추가할 기능", 2)
korean_with_translation(
    "앞으로는 이메일 알림 기능, 전자서명 연동, 그리고 AI 기반 계약서 검토 기능을 추가하고 싶습니다.",
    "In the future, I’d like to add email notifications, e-signature integration, and AI-based contract review."
)
bullet("Email notifications when a contract is sent or signed")
bullet("DocuSign or equivalent e-signature integration")
bullet("AI-powered contract review — flag unusual or non-standard clauses automatically")
bullet("Fund-level statistics dashboard (total invested, portfolio breakdown)")
bullet("Multiple interested party support per deal")
bullet("Contract version control — track revision history")
bullet("Admin UI for template editing — no more XML editing")

# ═══════════════════════════════════════════════════════════════════════════════
# SLIDE 16 — 라이브 데모
# ═══════════════════════════════════════════════════════════════════════════════
slide_label(16, "라이브 데모", "약 4~5분")
korean_with_translation(
    "이제 실제 웹사이트를 직접 시연해 보겠습니다.",
    "Now I will demonstrate the actual website directly."
)

heading("데모 순서 (Demo Flow)", 2)
bullet("0.  Open browser, go to sparklabs-doc-tracker.vercel.app  (already logged in)")
bullet("1.  메인 대시보드 — show KPI tiles, progress bars, deal list")
note("'메인 대시보드에서 현재 진행 중인 딕들을 확인할 수 있습니다.'")
note("  ↳ EN: 'You can see all currently active deals on the main dashboard.'")
bullet("2.  액션 센터 — point out the urgency items, explain how they’re ranked")
bullet("3.  서류 수집 탭 — show drag-and-drop, checklist auto-check")
note("'서류를 업로드하면 체크리스트가 자동으로 체크됩니다.'")
note("  ↳ EN: 'When a document is uploaded, the checklist item is automatically checked off.'")
bullet("4.  RCPS 계약서 탭 — type a company name, watch it auto-fill, hit Download")
note("'RCPS 계약서 작성 탭입니다. 정보를 입력하면 자동으로 계약서가 완성됩니다.'")
note("  ↳ EN: 'This is the RCPS agreement tab. When you enter information, the agreement is automatically completed.'")
bullet("     → Fill in date + investment amount → click Download → open the .docx to show it’s complete")
bullet("5.  AI 챗봇 — ask a process question, show the context-aware Korean/English answer")
note("'평상시처럼 질문을 하면 AI가 즈시 답변합니다.'")
note("  ↳ EN: 'You can ask a question just like you normally would, and the AI answers immediately.'")

korean_with_translation(
    "이상으로 발표를 마치겠습니다. 질문 있으시면 말씨해 주세요. 감사합니다.",
    "This concludes my presentation. Please feel free to ask any questions. Thank you."
)

# ═══════════════════════════════════════════════════════════════════════════════
# Q&A PREP
# ═══════════════════════════════════════════════════════════════════════════════
divider()
heading("예상 질문 & 답변  (Q&A Prep)")

qa_pairs = [
    (
        "보안은 어떻게 처리했나요?  /  How did you handle security?",
        "웹사이트 접근은 스파크랙 계정(Google OAuth)으로만 가능합니다. "
        "파일은 Vercel Blob에 비공개로 저장되며, Drive 폴더도 내부 팀에만 공유됩니다.\n"
        "↳ EN: Access requires a SparkLabs Google account (OAuth). Files are stored privately on Vercel Blob, "
        "and Drive folders are only shared with internal team members."
    ),
    (
        "Gemini API 비용은 얼마나 되나요?  /  How much does the Gemini API cost?",
        "연간 약 $0.09 — 커피 한 잔보다 저렴합니다. Google의 무료 티어 범위 안에서 운영됩니다.\n"
        "↳ EN: About $0.09 per year — less than a cup of coffee. It runs comfortably within Google’s free tier."
    ),
    (
        "계약서 템플릿을 수정하려면?  /  How do you update the agreement templates?",
        "현재는 DOCX 파일의 XML을 직접 수정해야 합니다. "
        "향후에는 웹 UI에서 관리자가 직접 수정할 수 있도록 개선할 계획입니다.\n"
        "↳ EN: Currently the DOCX XML must be edited directly. "
        "In the future I plan to build an admin UI so templates can be managed without code changes."
    ),
    (
        "AI를 어떻게 활용했나요?  /  How did you use AI?",
        "두 가지입니다. 앱 내부에서는 Google Gemini API로 6가지 기능을 구현했고, "
        "개발 과정에서는 Claude와 Claude Code로 코드를 작성했습니다. "
        "프로그래밍 경험이 없었는데 AI 덕분에 개발할 수 있었습니다.\n"
        "↳ EN: Two ways: inside the app, Google Gemini API for 6 different features; "
        "for development, Claude and Claude Code to write the actual code. "
        "I had no prior programming experience, and AI made this possible."
    ),
    (
        "얼마나 걸려서 만들었나요?  /  How long did it take to build?",
        "약 [N]주 동안 개발했습니다. "
        "기본 기능부터 시작해 계약서 자동화, Drive 연동 등을 단계적으로 추가했습니다.\n"
        "↳ EN: About [N] weeks. I started with the core features and added automation, Drive integration, "
        "AI features, and the bilingual toggle step by step."
    ),
]

for q, a in qa_pairs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    qr = p.add_run(f"Q:  {q}")
    style_run(qr, 11, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(4)
    ar = p2.add_run(f"A:  {a}")
    style_run(ar, 11, color=BLACK)

# ── Save ─────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), "script-v2.docx")
doc.save(out)
print(f"Saved: {out}")
